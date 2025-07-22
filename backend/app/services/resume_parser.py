import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union
import re
import pymupdf  # PyMuPDF
import pdfplumber
from docx import Document
from pydantic import BaseModel

class ParsedResumeData(BaseModel):
    """Structured resume data model"""
    personal_info: Dict[str, Optional[str]] = {}
    experience: List[Dict[str, str]] = []
    education: List[Dict[str, str]] = []
    skills: List[str] = []
    projects: List[Dict[str, str]] = []
    certifications: List[str] = []
    languages: List[str] = []

class ResumeParserService:
    """Service for parsing resume files and extracting structured data"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="resume_parser_")
    
    def validate_file(self, file_path: str, file_size: int) -> tuple[bool, str]:
        """Validate file type and size"""
        path = Path(file_path)
        
        # Check file extension
        if path.suffix.lower() not in self.ALLOWED_EXTENSIONS:
            return False, f"Unsupported file type. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}"
        
        # Check file size
        if file_size > self.MAX_FILE_SIZE:
            return False, f"File too large. Max size: {self.MAX_FILE_SIZE / (1024*1024):.1f}MB"
        
        return True, "Valid file"
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Method 1: Try PyMuPDF first
        try:
            doc = pymupdf.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            if text.strip():
                return text
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
        
        # Method 2: Fallback to pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        
        raise ValueError("Could not extract text from PDF")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise ValueError(f"Could not read TXT file: {e}")
        except Exception as e:
            raise ValueError(f"Could not extract text from TXT: {e}")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text based on file type"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def extract_personal_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract personal information from resume text"""
        personal_info = {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "location": None
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            personal_info["email"] = email_match.group()
        
        # Extract phone number
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # US format
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',   # (123) 456-7890
            r'\+\d{1,3}[-.\s]?\d{1,14}',      # International
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                personal_info["phone"] = phone_match.group()
                break
        
        # Extract LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            personal_info["linkedin"] = linkedin_match.group()
        
        # Extract name (first few lines, excluding email/phone)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:5]:  # Check first 5 lines
            # Skip lines with email, phone, or common headers
            if any(skip in line.lower() for skip in ['email', 'phone', 'resume', 'cv', '@']):
                continue
            # Look for likely name (2-4 words, mostly letters)
            if re.match(r'^[A-Za-z\s]{2,50}$', line) and 2 <= len(line.split()) <= 4:
                personal_info["name"] = line
                break
        
        return personal_info
    
    def extract_sections(self, text: str) -> Dict[str, List[str]]:
        """Extract different sections from resume text"""
        sections = {
            "experience": [],
            "education": [],
            "skills": [],
            "projects": [],
            "certifications": []
        }
        
        # Define section headers
        section_patterns = {
            "experience": r'(work\s+experience|experience|employment|professional\s+experience)',
            "education": r'(education|academic|qualifications)',
            "skills": r'(skills|technical\s+skills|core\s+competencies|technologies)',
            "projects": r'(projects|personal\s+projects|side\s+projects)',
            "certifications": r'(certifications|certificates|licenses)'
        }
        
        # Split text into lines and process
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            section_found = None
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    section_found = section_name
                    break
            
            if section_found:
                # Save previous section content
                if current_section and current_content:
                    sections[current_section] = current_content
                
                # Start new section
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = current_content
        
        return sections
    
    def parse_resume(self, file_path: str, file_size: int) -> ParsedResumeData:
        """Main method to parse resume and return structured data"""
        # Validate file
        is_valid, message = self.validate_file(file_path, file_size)
        if not is_valid:
            raise ValueError(message)
        
        # Extract raw text
        raw_text = self.extract_text(file_path)
        if not raw_text.strip():
            raise ValueError("No text could be extracted from the file")
        
        # Extract personal information
        personal_info = self.extract_personal_info(raw_text)
        
        # Extract sections
        sections = self.extract_sections(raw_text)
        
        # Create structured data
        parsed_data = ParsedResumeData(
            personal_info=personal_info,
            experience=[{"content": "\n".join(sections["experience"])}] if sections["experience"] else [],
            education=[{"content": "\n".join(sections["education"])}] if sections["education"] else [],
            skills=sections["skills"],
            projects=[{"content": "\n".join(sections["projects"])}] if sections["projects"] else [],
            certifications=sections["certifications"]
        )
        
        return parsed_data
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except Exception:
            pass